"""
PaperGen_Pro - Word 文档生成模块

使用 python-docx 将生成的论文内容（大纲、章节正文、图片）
组装为格式化的 Word 文档。
"""
import os
import re
import hashlib
import urllib.request
import urllib.parse

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config


def generate_docx(
    outline: dict,
    sections_content: dict,
    images_data: list,
    used_references: list = None,
) -> str:
    """
    生成 Word 文档。

    Args:
        outline: 论文大纲 {"title": "...", "sections": [...]}.
        sections_content: 各章节正文 {"1. 引言": "正文...", ...}.
        images_data: 提取的图片信息列表。

    Returns:
        str: 生成的 Word 文件路径。
    """
    # 确保输出目录存在
    os.makedirs(config.TEMP_OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # ===== 设置标题 =====
    title = outline.get("title", "未命名论文")
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 摘要 =====
    abstract_points = outline.get("abstract_points", [])
    if abstract_points:
        doc.add_heading("摘要", level=1)
        abstract_text = "；".join(abstract_points) + "。"
        doc.add_paragraph(abstract_text)

    # ===== 各章节正文 =====
    sections = outline.get("sections", [])
    for section in sections:
        heading = section.get("heading", "")
        points = section.get("points", [])

        # 写入章节标题
        doc.add_heading(heading, level=1)

        # 写入章节正文
        content = sections_content.get(heading, "")
        if content:
            pattern = r'\[INSERT_IMG_([^\]]+)\]'
            parts = re.split(pattern, content)
            img_dict = {img.get("id"): img for img in images_data if img.get("id")}
            
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    text_chunk = part.strip()
                    if not text_chunk:
                        continue
                    paragraphs = text_chunk.split("\n\n")
                    for para_text in paragraphs:
                        para_text = para_text.strip()
                        if not para_text:
                            continue
                        if para_text.startswith("### "):
                            sub_heading = para_text.lstrip("#").strip()
                            doc.add_heading(sub_heading, level=2)
                        elif para_text.startswith("## "):
                            continue
                        else:
                            _add_paragraph_with_math(doc, para_text, used_references)
                else:
                    img_id = part.strip()
                    if img_id in img_dict:
                        img_info = img_dict[img_id]
                        _insert_image(doc, img_info)
                        img_info["_inserted"] = True
                    else:
                        print(f"[DocWriter] WARNING: AI referenced missing image ID: {img_id}")
        else:
            doc.add_paragraph(f"（{heading} 的正文尚未生成）")

        # 废弃传统的硬插逻辑，完全由 AI 的占位符决定图片位置

    # ===== 附录：所有图片汇总 =====
    remaining_images = [
        img for img in images_data
        if not img.get("_inserted", False)
    ]
    if remaining_images:
        doc.add_heading("附录：图片汇总", level=1)
        for img_info in remaining_images:
            _insert_image(doc, img_info)

    # ===== 附录：参考文献 =====
    if used_references:
        doc.add_heading("参考文献 (References)", level=1)
        # 用带编号的列表排列
        for i, ref in enumerate(used_references, 1):
            p = doc.add_paragraph(f"[{i}] {ref.get('text', '')}")
            # 设置下悬挂缩进等样式
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-20)

    # ===== 保存文件 =====
    output_path = os.path.join(config.TEMP_OUTPUT_DIR, "paper_output.docx")
    doc.save(output_path)

    print(f"[DocWriter] Word 文档已生成: {output_path}")
    return output_path


def _insert_image(doc: Document, img_info: dict) -> None:
    """
    向 Word 文档中插入单张图片及其图注。
    """
    img_path = img_info.get("path", "")
    caption = img_info.get("caption_context", "")

    if not os.path.exists(img_path):
        print(f"[DocWriter] 图片文件不存在，跳过: {img_path}")
        return

    try:
        doc.add_picture(img_path, width=Inches(5.0))

        # 添加图注（去除冗长的来源页码记录，只保留最核心的图名或简述）
        caption_text = caption.strip() if caption and len(caption) < 200 else "Figure"

        caption_para = doc.add_paragraph(caption_text)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置图注字体为小号斜体
        for run in caption_para.runs:
            run.font.size = Pt(9)
            run.font.italic = True

        print(f"[DocWriter] 插入图片: {os.path.basename(img_path)}")

    except Exception as e:
        print(f"[DocWriter] 插入图片失败: {img_path}, 错误: {e}")

def _render_latex_to_image(latex_str: str) -> str:
    """调用外部 API 将 LaTeX 渲染为 PNG 图片并缓存"""
    try:
        encoded_eq = urllib.parse.quote(latex_str.strip())
        
        hasher = hashlib.md5()
        hasher.update(latex_str.encode("utf-8"))
        cache_key = hasher.hexdigest()
        
        output_path = os.path.join(config.TEMP_DIR, f"math_{cache_key}.png")
        if os.path.exists(output_path):
            return output_path
            
        url = f"https://latex.codecogs.com/png.image?\\dpi{{300}}\\bg{{white}}%20{encoded_eq}"
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=10) as response:
            with open(output_path, 'wb') as f:
                f.write(response.read())
        return output_path
    except Exception as e:
        print(f"[DocWriter] 公式渲染失败 '{latex_str[:20]}': {e}")
        return ""

def _add_paragraph_with_math(doc: Document, text: str, used_references: list = None) -> None:
    """
    向文档添加段落，自动解析：
    1. $...$ (行内) 和 $$...$$ (居中块) 公式，并将其渲染为图片。
    2. [REF_xxx] 参考文献占位符，转换为上标数字如 [1]。
    """
    p = doc.add_paragraph()
    
    # 扩展正则，把 [REF_xxx] 也单独捕获出来保留
    # (?s) 即 re.DOTALL，允许 . 匹配包含换行符在内的多行公式
    pattern = r'(?s)(\$\$.+?\$\$|\$.+?\$|\[REF_[^\]]+\])'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith("$$") and part.endswith("$$"):
            # Block math
            math_str = part[2:-2].replace("\n", " ").strip()
            img_path = _render_latex_to_image(math_str)
            if img_path:
                run = p.add_run()
                run.add_picture(img_path, height=Pt(16))
            else:
                p.add_run(part)
        elif part.startswith("$") and part.endswith("$"):
            # Inline math
            math_str = part[1:-1].replace("\n", " ").strip()
            img_path = _render_latex_to_image(math_str)
            if img_path:
                run = p.add_run()
                run.add_picture(img_path, height=Pt(11))
            else:
                p.add_run(part)
        elif part.startswith("[REF_") and part.endswith("]"):
            # Citation placeholder
            if used_references is not None:
                ref_id = part[5:-1]
                # 去查找它是第几个被使用的（1-based index）
                try:
                    index = next(i for i, r in enumerate(used_references) if r["id"] == ref_id) + 1
                    run = p.add_run(f"[{index}]")
                    run.font.superscript = True
                except StopIteration:
                    # 找不到这篇文献（理论上不会走到这里）
                    p.add_run(part)
            else:
                p.add_run(part)
        else:
            p.add_run(part)
